import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    # Customize heading style if needed
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        # docx doesn't fully support east asia fonts via simple api, but this works partially
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return heading

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_image(doc, img_path, width=Inches(6)):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=width)
        # add caption
        caption = doc.add_paragraph(os.path.basename(img_path).split('.')[0])
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.style = 'Caption'
    else:
        doc.add_paragraph(f"[Image not found: {img_path}]")

from docx.oxml.ns import qn

def main():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Title
    title = doc.add_heading('基于小红书平台的哈尔滨冬季冰雪旅游舆情与打卡空间特征分析', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Author / Date
    doc.add_paragraph('\n')
    
    # 摘要
    add_heading(doc, '摘要', 1)
    add_paragraph(doc, '随着冰雪旅游热度的持续攀升，哈尔滨作为中国冰雪旅游的核心目的地，其服务质量和空间承载力面临着巨大考验。本研究基于小红书平台获取了超过5000条高质量的游客打卡笔记与评论数据。综合运用大语言模型（LLM）自然语言处理技术、情感分析以及空间标签归一化方法，对哈尔滨冬季旅游的“网红打卡地空间结构特征”与“游客真实情绪及痛点”进行了全景式的大数据剖析。')
    add_paragraph(doc, '研究发现，哈尔滨的网红打卡地在空间与类型上高度向核心景区（冰雪大世界）与历史文化街区（中央大街、索菲亚教堂）集中。在情绪特征上，总体满意度虽然较高，但依然暴露出明显的结构性短板，尤其是“交通出行困难”、“防寒保暖设施不足”以及“部分场景隐性物价过高”三大核心痛点。本报告通过量化热力图和象限图，明确指出了高危预警区域，为政府和相关企业在未来的设施优化与资源配置提供了极具价值的数据支撑和决策依据。')

    # 第一章
    add_heading(doc, '一、 哈尔滨网红打卡地空间与分类研究', 1)
    
    add_heading(doc, '1.1 高频打卡地分布特征', 2)
    add_paragraph(doc, '通过对小红书笔记的地点标签进行清洗和归一化处理，统计得出了哈尔滨冬季最高频打卡的POI（兴趣点）列表。数据显示，冰雪大世界、中央大街和索菲亚教堂构成了绝对的“三巨头”引流核心。这也印证了冰雪体验与俄式风情是哈尔滨冬季旅游的两大基石。')
    add_image(doc, r'analysis_outputs_phase4\figures\xhs_checkin_top_poi_bar.png')
    
    add_heading(doc, '1.2 打卡地空间类型结构', 2)
    add_paragraph(doc, '将打卡地点划分为地标建筑、冰雪景观、特色餐饮、商业街区等类型后，发现在哈尔滨的打卡网络中，“地标建筑”和“冰雪景观”占据主导地位，而“特色餐饮”也展现出了极高的曝光量，说明美食打卡已经成为游客行程中不可或缺的一环。')
    add_image(doc, r'analysis_outputs_phase4\figures\xhs_checkin_poi_type_bar.png')

    add_heading(doc, '1.3 社交媒体内容表达方式分析', 2)
    add_paragraph(doc, '不同类型的打卡地在小红书上呈现出截然不同的内容生态。通过构建“打卡地类型 × 内容表达”热力图可以看出：在冰雪景观类地点，游客主要发布“经验攻略”与“避坑指南”；而在特色餐饮和地标建筑处，则更多以“纯展示（出片/打卡）”和“体验分享”为主。这提示旅游管理者，对于冰雪景区，行前的信息披露和攻略引导尤为重要。')
    add_image(doc, r'analysis_outputs_phase4\figures\xhs_checkin_poi_type_content_heatmap.png')

    add_heading(doc, '1.4 打卡地热度与体验风险象限分析', 2)
    add_paragraph(doc, '将各POI的“打卡热度”与“负面情绪比例”结合，构建了四象限图。第一象限（高热度、高负面）属于“网红高危区”，主要集中在冰雪大世界等核心景区，这些地方虽然流量极高，但由于排队、寒冷和交通问题，极易引发舆情反噬。第四象限（高热度、低负面）则是“优质核心区”，如索菲亚教堂，出片率高且体验相对稳定。')
    add_image(doc, r'analysis_outputs_phase4\figures\xhs_checkin_heat_sentiment_quadrant.png')

    # 第二章
    add_heading(doc, '二、 旅游舆情与游客情感分析', 1)

    add_heading(doc, '2.1 核心体验维度与情绪结构', 2)
    add_paragraph(doc, '将游客评论拆解为多个体验维度（如景观打卡、行程规划、餐饮消费等），发现整体以正面和中性情绪为主，体现了“讨好型市格”带来的良好口碑基础。但在“交通出行”和“服务与设施”维度，负面情绪占比明显异常，这是本轮分析的重点突破口。')
    add_image(doc, r'analysis_outputs\figures\aspect_sentiment_stacked_bar.png')

    add_heading(doc, '2.2 核心负面痛点精准识别', 2)
    add_paragraph(doc, '运用帕累托法则（80/20定律）提取核心痛点，数据显示：排队拥挤、交通不便（打车难、黑车）、防寒设施不足是构成游客负面体验的“三座大山”。结合热力图可见，交通痛点在“行程规划”中爆发，而排队痛点则深度绑定“景区游玩”维度。')
    add_image(doc, r'analysis_outputs\figures\painpoint_pareto.png')
    add_image(doc, r'analysis_outputs\figures\aspect_painpoint_heatmap_count.png')

    add_heading(doc, '2.3 行前信息服务缺口分析', 2)
    add_paragraph(doc, '通过挖掘提问类笔记，我们梳理出了游客最迫切需要的行前信息服务。研究发现，“穿搭/防寒建议”和“行程路线规划”是绝对的咨询热点。这表明哈尔滨作为极寒目的地，其信息壁垒依然较高，官方在冬季旅游前期需加大科普性和实操性攻略的投放力度。')
    add_image(doc, r'analysis_outputs_phase4\figures\consultation_gap_heatmap.png')

    add_heading(doc, '2.4 主次负面情绪的隐性风险', 2)
    add_paragraph(doc, '对负面舆情进行深度切片，发现主维度与次维度的交织规律。例如，在“景区游玩”不佳的表象下，往往隐藏着“基础设施差”的隐性抱怨；在“交通出行”不畅的背后，则伴随着“服务态度恶劣”的次生情绪。这提醒管理者，解决痛点需寻根溯源，不能仅停留在表层。')
    add_image(doc, r'analysis_outputs_phase4\figures\main_secondary_negative_heatmap.png')

    # 第三章
    add_heading(doc, '三、 核心结论与服务设施优化策略', 1)
    
    add_heading(doc, '3.1 空间布局优化：缓解单极化承载压力', 2)
    add_paragraph(doc, '数据表明冰雪大世界和中央大街承受了绝大部分的打卡热度与客流压力。建议在未来规划中，沿松花江两岸挖掘和培育“平替”型冰雪游乐节点，利用数字地图进行热力疏导，减轻单点交通与餐饮的过载现象。')

    add_heading(doc, '3.2 交通与御寒设施双管齐下', 2)
    add_paragraph(doc, '针对小红书爆发的“交通打车难”和“排队冻透”两大痛点，建议在冰雪大世界、雪乡等核心景区增设大容量、高频次的“温暖接驳专线”；同时在排队密集区增设临时保暖驿站或暖风通道，从根本上阻断恶劣天气带来的负面情绪蔓延。')

    add_heading(doc, '3.3 提升行前数字化服务体验', 2)
    add_paragraph(doc, '响应游客强烈的“穿搭”与“路线”咨询缺口，官方文旅部门可联合社交平台，推出权威的“哈尔滨冬季生存指南”小程序，集成实时温度、排队时长预警、防寒装备租赁指南等功能，实现从“被动应对”向“主动服务”的跨越。')

    # Save
    doc.save('哈尔滨冬季旅游小红书全景分析报告.docx')
    print("Report generated successfully at '哈尔滨冬季旅游小红书全景分析报告.docx'.")

if __name__ == '__main__':
    main()
